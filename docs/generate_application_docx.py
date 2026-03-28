#!/usr/bin/env python3
"""
Generate application documentation in DOCX from Markdown sources.

Run from project root:
  python3 docs/generate_application_docx.py

Requires: pip install python-docx
Output: docs/Application_Complete_Documentation.docx

Uses APPLICATION_GUIDE.md, BUSINESS_GUIDE.md, and TECHNICAL_REFERENCE.md
as the single source of truth so the DOCX stays in sync with the docs.
"""
import os
import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
sys.path.insert(0, str(ROOT))

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print("Install python-docx: pip install python-docx")
    sys.exit(1)


# ---- Helpers ----

def add_heading(doc, text, level=1):
    text = _strip_md_inline(text)
    return doc.add_heading(text, level=min(level, 3))


def add_para(doc, text, bold=False):
    text = _strip_md_inline(text)
    p = doc.add_paragraph()
    if bold:
        p.add_run(text).bold = True
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(_strip_md_inline(item), style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(_strip_md_inline(item), style="List Number")


def add_table_from_rows(doc, headers, rows, col_widths=None):
    ncols = len(headers)
    nrows = 1 + len(rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        if j < ncols:
            table.rows[0].cells[j].text = _strip_md_inline(str(h))
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < ncols and i + 1 < nrows:
                table.rows[i + 1].cells[j].text = _strip_md_inline(str(cell))
    if col_widths:
        for row in table.rows:
            for j, w in enumerate(col_widths):
                if j < len(row.cells):
                    try:
                        row.cells[j].width = Inches(w)
                    except Exception:
                        pass
    return table


def _strip_md_inline(text):
    """Remove markdown bold/link for plain DOCX (keep text)."""
    if not text or not isinstance(text, str):
        return str(text)
    # **bold** / *italic* -> content only
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    # [label](url) -> label
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


# ---- Markdown to DOCX ----

def _parse_md_table(lines, start_i):
    """Parse a markdown table starting at start_i. Returns (rows, next_line_index)."""
    rows = []
    i = start_i
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            return rows, i
        if not line.strip().startswith("|"):
            return rows, i
        cells = [c.strip() for c in line.split("|") if c.strip() or line.count("|") > 1]
        if not cells:
            i += 1
            continue
        # Skip separator row (|---|---|)
        if re.match(r"^[\s\-:]+$", "".join(cells)):
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i


def _is_list_item(line):
    s = line.lstrip()
    if re.match(r"^[\-\*]\s+", s) or re.match(r"^\d+\.\s+", s):
        return True
    return False


def _heading_level(line):
    m = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
    if m:
        return len(m.group(1)), m.group(2).strip()
    return 0, None


def append_markdown_to_document(doc, md_path):
    """Append content from a Markdown file to the document. Handles # headings, paragraphs, bullets, numbered lists, tables."""
    path = Path(md_path)
    if not path.is_absolute():
        path = DOCS_DIR / path
    if not path.exists():
        doc.add_paragraph(f"[Markdown file not found: {path}]")
        return
    text = path.read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # Horizontal rule / empty
        if stripped in ("---", "***", "") or not stripped:
            i += 1
            continue
        # Headings
        level, title = _heading_level(line)
        if title:
            add_heading(doc, title, level=level)
            i += 1
            continue
        # Table
        if stripped.startswith("|"):
            table_rows, next_i = _parse_md_table(lines, i)
            i = next_i
            if table_rows:
                headers = table_rows[0]
                data_rows = table_rows[1:]
                add_table_from_rows(doc, headers, data_rows)
            continue
        # Bullet or numbered list
        if _is_list_item(line):
            list_start = i
            list_items = []
            while i < len(lines) and _is_list_item(lines[i]):
                s = lines[i].lstrip()
                s = re.sub(r"^[\-\*]\s+", "", s, count=1)
                s = re.sub(r"^\d+\.\s+", "", s, count=1)
                list_items.append(s.strip())
                i += 1
            if list_items:
                use_numbered = bool(re.match(r"^\d+\.\s+", lines[list_start].lstrip()))
                style = "List Number" if use_numbered else "List Bullet"
                for item in list_items:
                    doc.add_paragraph(_strip_md_inline(item), style=style)
            continue
        # Paragraph (possibly multi-line)
        para_lines = []
        while i < len(lines):
            ln = lines[i]
            if ln.strip() == "" or _heading_level(ln)[1] or ln.strip().startswith("|") or _is_list_item(ln):
                break
            para_lines.append(ln.strip())
            i += 1
        if para_lines:
            add_para(doc, " ".join(para_lines))
    doc.add_paragraph()


def add_cover_and_toc(doc):
    doc.add_heading("Complete Application Documentation", 0)
    doc.add_paragraph("Multi-Industry SaaS Platform — Business & Technical Documentation")
    doc.add_paragraph()
    doc.add_paragraph(
        "This document is generated from the project Markdown docs: "
        "APPLICATION_GUIDE.md, BUSINESS_GUIDE.md, and TECHNICAL_REFERENCE.md. "
        "It covers business functionality, use cases with examples, technical architecture, "
        "API reference, and operational procedures."
    )
    doc.add_paragraph()
    doc.add_paragraph("Contents:")
    add_bullets(doc, [
        "Part 1 — Application Guide (overview, architecture, use cases by industry, demo)",
        "Part 2 — Business Guide (module-by-module, examples, use case diagrams, feature matrix)",
        "Part 3 — Technical Reference (APIs, fields, functions, sequence diagrams)",
    ])
    doc.add_page_break()


def add_demo_operations(doc):
    """Add Part 4: Demo tenants and operations (concise)."""
    add_heading(doc, "Part 4 — Demo Tenants and Operations", level=1)
    add_heading(doc, "4.1 Demo tenant naming", level=2)
    add_para(doc, "Demo tenants: ss_business_salon, ss_business_clinic, ss_business_gym, ss_business_school, ss_business_store, ss_business_camp, ss_business_car_showroom.")
    add_heading(doc, "4.2 Scripts and commands", level=2)
    add_para(doc, "Seed: python scripts/run_seed_domain.py --domain salon")
    add_para(doc, "Delete one: python scripts/run_delete_domain.py --domain salon")
    add_para(doc, "Delete all demo: python scripts/run_delete_all_demo.py")
    add_para(doc, "Explore: python scripts/explore_all_modules.py")
    add_heading(doc, "4.3 Deployment summary", level=2)
    add_para(doc, "API: uvicorn app.main:create_app --factory --reload --port 8000. Set MONGO_URI, JWT_SECRET; optional REDIS_URI, CORS_ORIGINS.")
    add_para(doc, "Admin UI: cd admin_ui && npm run dev (VITE_API_BASE=http://localhost:8000/v1). Production: npm run build, serve dist/.")
    add_para(doc, "Super Admin: scripts/create_super_admin.py or BOOT_SUPER_ADMIN_EMAIL / BOOT_SUPER_ADMIN_PASSWORD in .env.")
    doc.add_page_break()


def add_appendix(doc):
    add_heading(doc, "Appendix — Document references", level=1)
    add_bullets(doc, [
        "DOCUMENTATION_INDEX.md — Index of business vs technical docs",
        "APPLICATION_GUIDE.md — Architecture, RBAC, use cases by industry",
        "BUSINESS_GUIDE.md — Module-by-module business view, examples, use case diagrams",
        "TECHNICAL_REFERENCE.md — APIs, fields, functions, sequence diagrams",
        "AI_CAPABILITIES.md — AI behaviour and config",
        "DEPLOYMENT.md — Deployment",
        "whatsapp-workflow.md, whatsapp-workflow-actions.md — WhatsApp",
        "scripts/README.md — Seed/delete scripts",
    ])
    doc.add_paragraph("End of document.")


def main():
    doc = Document()
    add_cover_and_toc(doc)

    # Part 1: Application Guide
    add_heading(doc, "Part 1 — Application Guide", level=1)
    append_markdown_to_document(doc, DOCS_DIR / "APPLICATION_GUIDE.md")
    doc.add_page_break()

    # Part 2: Business Guide
    add_heading(doc, "Part 2 — Business Guide", level=1)
    append_markdown_to_document(doc, DOCS_DIR / "BUSINESS_GUIDE.md")
    doc.add_page_break()

    # Part 3: Technical Reference
    add_heading(doc, "Part 3 — Technical Reference", level=1)
    append_markdown_to_document(doc, DOCS_DIR / "TECHNICAL_REFERENCE.md")
    doc.add_page_break()

    # Part 4: Demo & operations
    add_demo_operations(doc)
    add_appendix(doc)

    out_path = DOCS_DIR / "Application_Complete_Documentation.docx"
    DOCS_DIR.mkdir(exist_ok=True)
    doc.save(str(out_path))
    print(f"Generated: {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
